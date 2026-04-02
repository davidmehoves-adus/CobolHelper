"""
md_to_docx.py — Convert markdown documentation to styled .docx files.

Usage:
    python md_to_docx.py <input.md> [output.docx]
    python md_to_docx.py --batch <directory>

If output path is omitted, writes to the same directory with .docx extension.
--batch converts all .md files in the given directory (skips dotfiles).
"""

import sys
import os
import argparse
import mistune
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)       # Headers, list bullets
MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x8A)      # Sub-headers
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)        # Body text
LIGHT_GRAY_BG = "E8ECF0"                       # Code block background
TABLE_HEADER_BG = "2D3748"                     # Table header (dark charcoal)
TABLE_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)   # Table header text
TABLE_ALT_ROW = "F0F3F7"                       # Alternating row tint
TABLE_BORDER = "B0B8C4"                        # Table border color
HR_COLOR = "1B3A5C"                            # Horizontal rule
BULLET_COLOR = DARK_BLUE


# ---------------------------------------------------------------------------
# Style setup
# ---------------------------------------------------------------------------

def _set_shading(p_element, color_hex):
    """Apply background shading to a paragraph XML element."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    p_element.get_or_add_pPr().append(shading)


def _set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color_hex, sz="4"):
    """Set all borders on a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)


def _add_horizontal_rule(doc):
    """Add a styled horizontal rule (thin dark blue line)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    bottom = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:color="{HR_COLOR}" w:space="1"/>'
        f'</w:pBdr>'
    )
    p_pr.append(bottom)
    return p


def setup_styles(doc):
    """Create custom styles in the document."""
    styles = doc.styles

    # -- Heading styles --
    for level, (size, color, bold, spacing_before) in enumerate([
        (22, DARK_BLUE,   True,  18),   # Heading 1
        (17, DARK_BLUE,   True,  14),   # Heading 2
        (14, MEDIUM_BLUE, True,  12),   # Heading 3
        (12, MEDIUM_BLUE, True,  10),   # Heading 4
    ], start=1):
        style_name = f"Heading {level}"
        if style_name in [s.name for s in styles]:
            style = styles[style_name]
        else:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        fmt = style.font
        fmt.size = Pt(size)
        fmt.color.rgb = color
        fmt.bold = bold
        fmt.name = "Calibri"
        style.paragraph_format.space_before = Pt(spacing_before)
        style.paragraph_format.space_after = Pt(4)

    # -- Normal body text --
    normal = styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK_GRAY
    normal.font.name = "Calibri"
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = Pt(15)

    # -- Code Block style (paragraph) --
    if "Code Block" not in [s.name for s in styles]:
        code_block = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_block.font.name = "Consolas"
        code_block.font.size = Pt(9)
        code_block.font.color.rgb = DARK_GRAY
        code_block.paragraph_format.space_before = Pt(0)
        code_block.paragraph_format.space_after = Pt(0)
        code_block.paragraph_format.line_spacing = Pt(13)
        code_block.paragraph_format.left_indent = Cm(0.5)


# ---------------------------------------------------------------------------
# Markdown AST -> docx rendering
# ---------------------------------------------------------------------------

class DocxRenderer:
    """Walk a mistune 2.x AST and produce a python-docx Document."""

    def __init__(self):
        self.doc = Document()
        setup_styles(self.doc)
        # Remove the default empty paragraph
        if self.doc.paragraphs:
            p = self.doc.paragraphs[0]._element
            p.getparent().remove(p)

    def render(self, tokens):
        for tok in tokens:
            self._render_token(tok)
        return self.doc

    # -- Token dispatch --

    def _render_token(self, tok):
        tp = tok["type"]
        dispatch = {
            "heading":        self._heading,
            "paragraph":      self._paragraph,
            "block_code":     self._block_code,
            "list":           self._list,
            "thematic_break": self._thematic_break,
            "table":          self._table,
            "blank_line":     self._blank_line,
            "block_quote":    self._block_quote,
        }
        handler = dispatch.get(tp)
        if handler:
            handler(tok)

    # -- Block-level handlers --

    def _heading(self, tok):
        level = tok.get("level", 1)
        style_name = f"Heading {min(level, 4)}"
        p = self.doc.add_paragraph(style=style_name)
        self._render_inline(p, tok.get("children", []))

    def _paragraph(self, tok):
        p = self.doc.add_paragraph()
        self._render_inline(p, tok.get("children", []))

    def _block_code(self, tok):
        raw = tok.get("text", tok.get("raw", ""))
        lines = raw.rstrip("\n").split("\n")
        for line in lines:
            p = self.doc.add_paragraph(style="Code Block")
            _set_shading(p._p, LIGHT_GRAY_BG)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line if line else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        # Spacer after code block
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(2)
        spacer.paragraph_format.space_after = Pt(2)

    def _list(self, tok, indent_level=0):
        ordered = tok.get("ordered", False)
        counter = tok.get("start", 1) or 1
        for item in tok.get("children", []):
            self._render_list_item(item, ordered, counter, indent_level)
            if ordered:
                counter += 1

    def _render_list_item(self, tok, ordered, counter, indent_level):
        children = tok.get("children", [])
        first = True
        for child in children:
            ctype = child["type"]
            if ctype == "list":
                self._list(child, indent_level + 1)
            elif ctype in ("paragraph", "block_text"):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.8)
                if first:
                    bullet = f"{counter}." if ordered else "\u2022"
                    run = p.add_run(f"{bullet}  ")
                    run.font.color.rgb = BULLET_COLOR
                    run.font.bold = True
                    run.font.size = Pt(10.5)
                    first = False
                self._render_inline(p, child.get("children", []))
            elif ctype == "block_code":
                self._block_code(child)
            else:
                self._render_token(child)

    def _thematic_break(self, tok):
        _add_horizontal_rule(self.doc)

    def _blank_line(self, tok):
        pass

    def _block_quote(self, tok):
        for child in tok.get("children", []):
            if child["type"] in ("paragraph", "block_text"):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0)
                self._render_inline(p, child.get("children", []))
                for run in p.runs:
                    run.font.italic = True
                    run.font.color.rgb = MEDIUM_BLUE
            else:
                self._render_token(child)

    # -- Table --

    def _table(self, tok):
        children = tok.get("children", [])

        # Collect header cells and body rows
        header_cells = []
        body_rows = []
        for section in children:
            stype = section["type"]
            if stype == "table_head":
                # table_head contains table_cell children directly
                header_cells = [
                    self._inline_text(cell.get("children", []))
                    for cell in section.get("children", [])
                    if cell["type"] == "table_cell"
                ]
            elif stype == "table_body":
                for row in section.get("children", []):
                    body_rows.append([
                        self._inline_text(cell.get("children", []))
                        for cell in row.get("children", [])
                        if cell["type"] == "table_cell"
                    ])

        if not header_cells:
            return

        num_cols = len(header_cells)
        num_rows = 1 + len(body_rows)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True

        # Header row
        for i, text in enumerate(header_cells):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.bold = True
            run.font.color.rgb = TABLE_HEADER_FG
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            _set_cell_shading(cell, TABLE_HEADER_BG)
            _set_cell_borders(cell, TABLE_BORDER)

        # Body rows
        for row_idx, row_data in enumerate(body_rows):
            for col_idx, text in enumerate(row_data):
                if col_idx >= num_cols:
                    break
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.color.rgb = DARK_GRAY
                _set_cell_borders(cell, TABLE_BORDER)
                if row_idx % 2 == 1:
                    _set_cell_shading(cell, TABLE_ALT_ROW)

        # Spacer after table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(4)
        spacer.paragraph_format.space_after = Pt(4)

    # -- Inline rendering --

    def _render_inline(self, paragraph, children):
        """Render inline tokens into an existing paragraph."""
        if not children:
            return
        for child in children:
            self._render_inline_token(paragraph, child)

    def _render_inline_token(self, paragraph, tok):
        tp = tok["type"]
        if tp == "text":
            text = tok.get("text", tok.get("raw", ""))
            if isinstance(text, str):
                paragraph.add_run(text)
            elif isinstance(text, list):
                self._render_inline(paragraph, text)
        elif tp == "codespan":
            text = tok.get("text", tok.get("raw", ""))
            run = paragraph.add_run(text if isinstance(text, str) else str(text))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_BLUE
        elif tp == "strong":
            runs_before = len(paragraph.runs)
            for ch in tok.get("children", []):
                self._render_inline_token(paragraph, ch)
            for run in paragraph.runs[runs_before:]:
                run.font.bold = True
        elif tp == "emphasis":
            runs_before = len(paragraph.runs)
            for ch in tok.get("children", []):
                self._render_inline_token(paragraph, ch)
            for run in paragraph.runs[runs_before:]:
                run.font.italic = True
        elif tp == "link":
            text = self._inline_text(tok.get("children", []))
            run = paragraph.add_run(text)
            run.font.color.rgb = MEDIUM_BLUE
            run.font.underline = True
        elif tp in ("softbreak", "linebreak"):
            paragraph.add_run("\n")
        else:
            # Fallback: try to extract any text content
            text = tok.get("text", tok.get("raw", ""))
            children = tok.get("children", None)
            if isinstance(text, str) and text:
                paragraph.add_run(text)
            elif isinstance(children, list):
                self._render_inline(paragraph, children)

    def _inline_text(self, children):
        """Extract plain text from inline tokens (for table cells, etc.)."""
        parts = []
        if not children:
            return ""
        for tok in children:
            tp = tok["type"]
            if tp == "text":
                text = tok.get("text", tok.get("raw", ""))
                if isinstance(text, str):
                    parts.append(text)
                elif isinstance(text, list):
                    parts.append(self._inline_text(text))
            elif tp == "codespan":
                text = tok.get("text", tok.get("raw", ""))
                parts.append(text if isinstance(text, str) else str(text))
            elif tp in ("strong", "emphasis", "link"):
                parts.append(self._inline_text(tok.get("children", [])))
            elif tp in ("softbreak", "linebreak"):
                parts.append(" ")
            else:
                text = tok.get("text", tok.get("raw", ""))
                if isinstance(text, str):
                    parts.append(text)
                children_inner = tok.get("children", None)
                if isinstance(children_inner, list):
                    parts.append(self._inline_text(children_inner))
        return "".join(parts)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def convert_file(md_path, docx_path=None):
    """Convert a single .md file to .docx. Returns the output path."""
    if docx_path is None:
        docx_path = os.path.splitext(md_path)[0] + ".docx"

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Parse markdown into AST with table plugin
    parser = mistune.create_markdown(renderer="ast", plugins=["table"])
    tokens = parser(md_text)

    # Render to docx
    renderer = DocxRenderer()
    doc = renderer.render(tokens)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.save(docx_path)
    return docx_path


def convert_batch(directory):
    """Convert all .md files in a directory. Returns list of output paths."""
    results = []
    for name in sorted(os.listdir(directory)):
        if name.startswith(".") or not name.endswith(".md"):
            continue
        md_path = os.path.join(directory, name)
        if os.path.isfile(md_path):
            out = convert_file(md_path)
            results.append(out)
            print(f"  Converted: {name} -> {os.path.basename(out)}")
    return results


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Convert markdown documentation to styled .docx files."
    )
    parser.add_argument("input", help="Markdown file path, or directory with --batch")
    parser.add_argument("output", nargs="?", help="Output .docx path (optional)")
    parser.add_argument("--batch", action="store_true",
                        help="Convert all .md files in the given directory")

    args = parser.parse_args()

    if args.batch:
        if not os.path.isdir(args.input):
            print(f"Error: {args.input} is not a directory", file=sys.stderr)
            sys.exit(1)
        results = convert_batch(args.input)
        print(f"\nConverted {len(results)} file(s).")
    else:
        if not os.path.isfile(args.input):
            print(f"Error: {args.input} not found", file=sys.stderr)
            sys.exit(1)
        out = convert_file(args.input, args.output)
        print(f"Converted: {out}")


if __name__ == "__main__":
    main()
